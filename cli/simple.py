#!/usr/bin/env python3
"""
Simple CLI - One-command interface for non-technical users.

Usage:
    lingframe analyze essay.pdf
    lingframe analyze essay.pdf --output report.html
    lingframe analyze essay.pdf --advanced

This module provides a streamlined interface that hides all technical
complexity behind a single command. It automatically:
1. Creates a temporary project structure
2. Atomizes the document
3. Runs the evaluation analysis
4. Generates a narrative report
5. Opens the report in the default browser
"""

from __future__ import annotations

import argparse
import json
import shutil
import sys
import tempfile
import webbrowser
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import uuid


def get_framework_root() -> Path:
    """Get the framework root directory."""
    return Path(__file__).resolve().parent.parent


def setup_framework():
    """Initialize framework and register components."""
    framework_root = get_framework_root()
    sys.path.insert(0, str(framework_root))

    from framework.core import registry
    from framework.domains import DOMAINS_DIR

    # Discover domains
    registry.discover_domains(DOMAINS_DIR)

    # Import analysis modules
    from framework.analysis import (
        SemanticAnalysis,
        TemporalAnalysis,
        SentimentAnalysis,
        EntityAnalysis,
        EvaluationAnalysis,
    )

    # Import visualization adapters
    from framework.visualization import (
        ForceGraphAdapter,
        SankeyAdapter,
        SentimentChartAdapter,
        EntityBrowserAdapter,
        EvaluationDashboardAdapter,
    )

    return registry


def detect_file_type(file_path: Path) -> str:
    """Detect the type of input file."""
    suffix = file_path.suffix.lower()
    type_map = {
        ".pdf": "pdf",
        ".txt": "text",
        ".md": "markdown",
        ".docx": "docx",
        ".doc": "doc",
    }
    return type_map.get(suffix, "unknown")


def extract_text(file_path: Path) -> str:
    """Extract text from various file formats."""
    file_type = detect_file_type(file_path)

    if file_type == "pdf":
        from framework.loaders import PDFLoader
        loader = PDFLoader()
        return loader.extract_text(file_path)

    elif file_type in ("text", "markdown"):
        return file_path.read_text(encoding="utf-8")

    elif file_type == "docx":
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX files. "
                "Install with: pip install python-docx"
            )

    else:
        # Try reading as plain text
        try:
            return file_path.read_text(encoding="utf-8")
        except Exception:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")


def create_temp_project(
    text: str,
    document_title: str,
    temp_dir: Path,
) -> Tuple[Dict[str, Any], Path]:
    """
    Create a temporary project structure for analysis.

    Returns:
        Tuple of (project config, project directory)
    """
    project_dir = temp_dir / "project"
    project_dir.mkdir(parents=True, exist_ok=True)

    # Create directory structure
    (project_dir / "docs").mkdir(exist_ok=True)
    (project_dir / "data" / "raw").mkdir(parents=True, exist_ok=True)
    (project_dir / "data" / "processed").mkdir(parents=True, exist_ok=True)

    # Save text to a file
    text_path = project_dir / "docs" / "document.txt"
    text_path.write_text(text, encoding="utf-8")

    # Create project config
    config = {
        "project": {
            "name": "temp-analysis",
            "version": "1.0.0",
        },
        "naming": {
            "strategy": "legacy",  # Simple naming for temp projects
        },
        "corpus": {
            "documents": [
                {
                    "source": "docs/document.txt",
                    "title": document_title,
                    "id": "doc-001",
                }
            ]
        },
        "domain": {
            "profile": "base",  # Default domain
        },
        "analysis": {
            "pipelines": [
                {"module": "evaluation"},
                {"module": "semantic"},
                {"module": "sentiment"},
            ]
        },
        "output": {
            "raw_dir": "data/raw",
            "processed_dir": "data/processed",
        }
    }

    # Write config
    try:
        import yaml
        config_path = project_dir / "project.yaml"
        with open(config_path, "w", encoding="utf-8") as f:
            yaml.dump(config, f, default_flow_style=False)
    except ImportError:
        # Fall back to JSON if yaml not available
        config_path = project_dir / "project.json"
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2)

    return config, project_dir


def run_analysis(
    text: str,
    document_title: str,
    registry,
    verbose: bool = False,
) -> Dict[str, Any]:
    """
    Run analysis on the text and return results.

    Returns:
        Dict containing evaluation and optional supplementary analyses
    """
    from framework.core import Atomizer, AtomizationSchema, Corpus
    from framework.core.ontology import Document, AtomLevel

    # Create atomization schema
    schema = AtomizationSchema.default()

    # Create atomizer and process text
    atomizer = Atomizer(schema)
    atomizer._reset_counters()

    if verbose:
        print("  Analyzing text structure...")

    # Create document and atomize text directly
    doc = Document(
        id="doc-001",
        source_path=Path("inline-text"),
        format="plain",
        title=document_title,
    )

    # Start atomization from first level
    first_level = schema.levels[0]
    doc.root_atoms = atomizer.atomize_text(text, first_level)

    corpus = Corpus(
        name=document_title,
        documents=[doc],
        schema=schema,
    )

    if verbose:
        from framework.core.ontology import AtomLevel
        print(f"  Found {corpus.count_atoms(AtomLevel.THEME)} themes, "
              f"{corpus.count_atoms(AtomLevel.PARAGRAPH)} paragraphs, "
              f"{corpus.count_atoms(AtomLevel.SENTENCE)} sentences")

    # Get domain profile
    domain = registry.get_domain("base")

    results = {}

    # Run evaluation analysis (primary)
    if verbose:
        print("  Running rhetorical evaluation...")

    eval_module = registry.create_analysis("evaluation")
    eval_output = eval_module.analyze(corpus, domain, {})
    results["evaluation"] = eval_output.to_dict()

    # Run supplementary analyses
    try:
        if verbose:
            print("  Analyzing themes and connections...")
        semantic_module = registry.create_analysis("semantic")
        semantic_output = semantic_module.analyze(corpus, domain, {})
        results["semantic"] = semantic_output.to_dict()
    except Exception as e:
        if verbose:
            print(f"  Warning: Semantic analysis skipped ({e})")

    try:
        if verbose:
            print("  Analyzing emotional tone...")
        sentiment_module = registry.create_analysis("sentiment")
        sentiment_output = sentiment_module.analyze(corpus, domain, {})
        results["sentiment"] = sentiment_output.to_dict()
    except Exception as e:
        if verbose:
            print(f"  Warning: Sentiment analysis skipped ({e})")

    return results


def generate_report(
    analysis_results: Dict[str, Any],
    document_title: str,
    output_path: Path,
    verbose: bool = False,
) -> Path:
    """
    Generate a narrative HTML report from analysis results.

    Returns:
        Path to the generated report
    """
    from framework.output import NarrativeReportGenerator

    if verbose:
        print("  Generating narrative report...")

    generator = NarrativeReportGenerator(include_icons=True, verbose=verbose)

    # Extract evaluation data (primary) and supplementary analyses
    evaluation_data = analysis_results.get("evaluation", {})
    additional = {
        k: v for k, v in analysis_results.items()
        if k != "evaluation"
    }

    # Generate report
    report = generator.generate(
        evaluation_data=evaluation_data,
        document_title=document_title,
        additional_analyses=additional,
    )

    # Render to HTML
    html_content = generator.to_html(report)

    # Write output
    output_path.write_text(html_content, encoding="utf-8")

    if verbose:
        print(f"  Report saved: {output_path}")

    return output_path


def cmd_analyze(args: argparse.Namespace) -> int:
    """
    Main analyze command - the one-command experience.

    Takes a document file and produces a narrative report.
    """
    input_file = Path(args.file).resolve()

    # Validate input file
    if not input_file.exists():
        print(f"Error: File not found: {input_file}")
        return 1

    # Determine document title
    document_title = args.title or input_file.stem.replace("-", " ").replace("_", " ").title()

    # Determine output path
    if args.output:
        output_path = Path(args.output).resolve()
    else:
        output_path = input_file.parent / f"{input_file.stem}_analysis.html"

    print(f"Analyzing: {input_file.name}")
    print("=" * 50)

    try:
        # Setup framework
        if args.verbose:
            print("  Initializing analysis engine...")
        registry = setup_framework()

        # Extract text from document
        if args.verbose:
            print("  Extracting text from document...")
        text = extract_text(input_file)

        if not text.strip():
            print("Error: No text found in document")
            return 1

        if args.verbose:
            word_count = len(text.split())
            print(f"  Extracted {word_count:,} words")

        # Run analysis
        results = run_analysis(
            text=text,
            document_title=document_title,
            registry=registry,
            verbose=args.verbose,
        )

        # Generate report
        report_path = generate_report(
            analysis_results=results,
            document_title=document_title,
            output_path=output_path,
            verbose=args.verbose,
        )

        print()
        print("Analysis complete!")
        print(f"Report: {report_path}")

        # Export raw JSON if requested
        if args.json:
            json_path = output_path.with_suffix(".json")
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            print(f"JSON data: {json_path}")

        # Open in browser unless --no-open
        if not args.no_open:
            if args.verbose:
                print("  Opening in browser...")
            webbrowser.open(f"file://{report_path}")

        return 0

    except ImportError as e:
        print(f"Error: Missing dependency - {e}")
        print("Try: pip install -r requirements.txt")
        return 1

    except Exception as e:
        print(f"Error: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        return 1


def cmd_quick(args: argparse.Namespace) -> int:
    """
    Quick analysis - minimal output, just the key insights.

    Prints a brief summary to the console instead of generating a full report.
    """
    input_file = Path(args.file).resolve()

    if not input_file.exists():
        print(f"Error: File not found: {input_file}")
        return 1

    document_title = args.title or input_file.stem.replace("-", " ").replace("_", " ").title()

    print(f"Quick Analysis: {document_title}")
    print("=" * 50)

    try:
        registry = setup_framework()
        text = extract_text(input_file)

        if not text.strip():
            print("Error: No text found")
            return 1

        # Run minimal analysis
        results = run_analysis(
            text=text,
            document_title=document_title,
            registry=registry,
            verbose=False,
        )

        # Print quick summary
        eval_data = results.get("evaluation", {})
        summary = eval_data.get("summary", {})

        overall_score = summary.get("overall_score", 0)
        phase_scores = summary.get("phase_scores", {})
        recommendations = summary.get("top_recommendations", [])

        print()
        # Scores from evaluation module are already 0-100
        print(f"Overall Score: {int(overall_score)}%")
        print()

        from framework.output.terminology import friendly

        if phase_scores:
            print("By Phase:")
            for phase, score in sorted(phase_scores.items(), key=lambda x: -x[1]):
                phase_name = friendly(phase)
                # Scores from evaluation module are already 0-100
                print(f"  {phase_name}: {int(score)}%")
            print()

        if recommendations:
            print("Top Recommendations:")
            for i, rec in enumerate(recommendations[:3], 1):
                # Truncate long recommendations
                if len(rec) > 80:
                    rec = rec[:77] + "..."
                print(f"  {i}. {rec}")

        print()
        print("Run 'lingframe analyze <file>' for a detailed report.")

        return 0

    except Exception as e:
        print(f"Error: {e}")
        return 1


def create_simple_parser() -> argparse.ArgumentParser:
    """Create the simplified argument parser."""
    parser = argparse.ArgumentParser(
        prog="lingframe",
        description="Rhetorical analysis for writers - no coding required",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  lingframe analyze essay.pdf           Analyze and open report in browser
  lingframe analyze essay.pdf -o out.html   Save report to specific location
  lingframe quick essay.pdf             Quick console summary (no file)
  lingframe analyze essay.pdf --json    Also export raw data as JSON

For advanced options, use the full CLI:
  lingframe run --project my-project --visualize
        """,
    )

    subparsers = parser.add_subparsers(dest="command", help="Commands")

    # analyze command (primary)
    analyze_parser = subparsers.add_parser(
        "analyze",
        help="Analyze a document and generate a report",
        description="Analyze a document's rhetorical structure and generate "
                    "an interactive HTML report.",
    )
    analyze_parser.add_argument(
        "file",
        help="Document to analyze (PDF, TXT, MD, DOCX)"
    )
    analyze_parser.add_argument(
        "-o", "--output",
        help="Output file path (default: <input>_analysis.html)"
    )
    analyze_parser.add_argument(
        "-t", "--title",
        help="Document title (default: derived from filename)"
    )
    analyze_parser.add_argument(
        "--json",
        action="store_true",
        help="Also export raw analysis data as JSON"
    )
    analyze_parser.add_argument(
        "--no-open",
        action="store_true",
        help="Don't open the report in browser"
    )
    analyze_parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Show detailed progress"
    )
    analyze_parser.add_argument(
        "--advanced",
        action="store_true",
        help="Show advanced options (for power users)"
    )
    analyze_parser.set_defaults(func=cmd_analyze)

    # quick command (console-only summary)
    quick_parser = subparsers.add_parser(
        "quick",
        help="Quick analysis - console summary only",
        description="Get a quick summary of your document's rhetorical "
                    "strengths and weaknesses without generating a full report.",
    )
    quick_parser.add_argument(
        "file",
        help="Document to analyze"
    )
    quick_parser.add_argument(
        "-t", "--title",
        help="Document title"
    )
    quick_parser.set_defaults(func=cmd_quick)

    return parser


def main_simple() -> int:
    """Entry point for simplified CLI."""
    parser = create_simple_parser()
    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return 0

    return args.func(args)


if __name__ == "__main__":
    sys.exit(main_simple())
